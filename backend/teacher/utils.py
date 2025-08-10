import logging
import os
import cloudinary
from uuid import uuid4
from fastapi import UploadFile
import cloudinary.uploader
from config import settings
import io
import mimetypes
import openpyxl
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

cloudinary.config(
    cloud_name=settings.cloud_name,
    api_key=settings.cloudinary_api_key,
    api_secret=settings.cloudinary_api_secret,
    secure=settings.cloud_secure,  # Usa HTTPS si es True
)


# -- Funciones de utilidad para maneja
# r archivos y Cloudinary
async def save_and_upload_file(
    file: UploadFile, upload_folder: str = "uploads", public_id: str = None
) -> str:
    """
    Guarda un archivo localmente, lo sube a Cloudinary y elimina el archivo local.
    :param file: UploadFile de FastAPI.
    :param upload_folder: Carpeta local temporal.
    :return: URL del archivo en Cloudinary.
    """
    if not os.path.exists(upload_folder):
        os.makedirs(upload_folder)

    # Nombre único para el archivo
    filename = f"{uuid4().hex}_{file.filename}"
    local_path = os.path.join(upload_folder, filename)

    # Guarda el archivo localmente
    with open(local_path, "wb") as destination:
        content = await file.read()  # Lee todo el archivo
        destination.write(content)

    # Sube el archivo a Cloudinary y elimina local
    try:
        result = cloudinary.uploader.upload(
            local_path, public_id=public_id, resource_type="auto"
        )
        url = result.get("secure_url")
    finally:
        if os.path.exists(local_path):
            os.remove(local_path)

    return url


async def update_file_on_cloudinary(file: UploadFile, public_id: str) -> str:
    """
    Actualiza un archivo existente en Cloudinary reemplazando el archivo con el mismo public_id.
    :param file: UploadFile de FastAPI.
    :param public_id: ID público del archivo en Cloudinary a actualizar.
    :return: URL del archivo actualizado en Cloudinary.
    """
    filename = f"{uuid4().hex}_{file.filename}"
    local_path = os.path.join("uploads", filename)

    if not os.path.exists("uploads"):
        os.makedirs("uploads")

    with open(local_path, "wb") as destination:
        content = await file.read()
        destination.write(content)

    try:
        result = cloudinary.uploader.upload(local_path, public_id=public_id)
        url = result.get("secure_url")
    finally:
        if os.path.exists(local_path):
            os.remove(local_path)

    return url


async def delete_file_from_cloudinary(
    public_id: str, resource_type: str = "image"
) -> None:
    """
    Elimina un archivo de Cloudinary usando su public_id.
    :param public_id: ID público del archivo en Cloudinary a eliminar.
    :param resource_type: Tipo de recurso ('image', 'video', 'raw', etc.). Por defecto 'image'.
    """
    try:
        # Intenta eliminar con el tipo especificado
        result = cloudinary.uploader.destroy(public_id, resource_type=resource_type)

        # Si no se encuentra con el tipo especificado, intenta con otros tipos comunes
        if result.get("result") == "not found" and resource_type == "image":
            logger.info(
                f"No se encontró como imagen, intentando con otros tipos para: {public_id}"
            )

            # Intenta con video
            result = cloudinary.uploader.destroy(public_id, resource_type="video")
            if result.get("result") == "not found":
                # Intenta con raw (documentos, etc.)
                result = cloudinary.uploader.destroy(public_id, resource_type="raw")

        logger.info(f"Resultado de eliminación de Cloudinary: {result}")

    except Exception as e:
        logger.error(f"Error al eliminar el archivo de Cloudinary: {e}")


async def delete_file_from_cloudinary_auto(public_id: str) -> bool:
    """
    Elimina un archivo de Cloudinary detectando automáticamente el tipo de recurso.
    :param public_id: ID público del archivo en Cloudinary a eliminar.
    :return: True si se eliminó exitosamente, False en caso contrario.
    """
    resource_types = ["image", "video", "raw"]

    for resource_type in resource_types:
        try:
            result = cloudinary.uploader.destroy(public_id, resource_type=resource_type)
            if result.get("result") == "ok":
                logger.info(
                    f"Archivo eliminado exitosamente como {resource_type}: {public_id}"
                )
                return True
        except Exception as e:
            logger.warning(f"Error al intentar eliminar como {resource_type}: {e}")
            continue

    logger.error(
        f"No se pudo eliminar el archivo con ningún tipo de recurso: {public_id}"
    )
    return False


# -- Otras funciones de utilidad --
def generate_profile_prefix(name: str, last_name: str) -> str:
    """
    ## Generar prefijo

    Permite generar un prefijo en base a el nombre y el apellido del estudiante.

    ### Parámetros:
    - `name(str):` Nombre del estudiante.
    - `last_name(str)`: Apellido del estudiante.

    ### Retornos:
    - `String`: Prefijo del nombre y el apellido del estudiante.
    """

    first_letter_name: str = name[0].upper()
    first_letter_last_name: str = last_name[0].upper()

    prefix: str = first_letter_name + first_letter_last_name

    return prefix


async def read_student_identification(file: UploadFile) -> list:
    """
    Lee y procesa números de identificación desde un archivo de texto, Excel (.xlsx) o Word (.docx).

    :param file: Archivo con números de identificación.
    :return: Lista de números de identificación únicos y válidos.
    :raises ValueError: Si el formato no es soportado o hay errores en el archivo.
    :raises HTTPException: Si hay problemas con el tamaño o contenido del archivo.
    """
    # Validaciones iniciales
    if not file.filename:
        raise ValueError("El archivo debe tener un nombre válido")

    # Limitar tamaño de archivo (5MB)
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    content = await file.read()

    if len(content) > MAX_FILE_SIZE:
        raise ValueError("El archivo es demasiado grande (máximo 5MB)")

    if len(content) == 0:
        raise ValueError("El archivo está vacío")

    filename = file.filename.lower()
    mime_type, _ = mimetypes.guess_type(filename)
    identification_numbers = set()  # Usar set para evitar duplicados

    try:
        if filename.endswith(".txt") or (mime_type and mime_type.startswith("text")):
            # Intentar múltiples encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    text_content = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("No se pudo decodificar el archivo de texto")

            lines = text_content.splitlines()
            for line in lines:
                line = line.strip()
                if line and _validate_identification_number(line):
                    identification_numbers.add(line)

        elif filename.endswith(".xlsx"):
            wb = openpyxl.load_workbook(
                io.BytesIO(content), read_only=True, data_only=True
            )
            ws = wb.active

            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        cell_str = str(cell).strip()
                        if _validate_identification_number(cell_str):
                            identification_numbers.add(cell_str)
            wb.close()

        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                line = para.text.strip()
                if line and _validate_identification_number(line):
                    identification_numbers.add(line)

            # También revisar tablas en el documento
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and _validate_identification_number(cell_text):
                            identification_numbers.add(cell_text)

        else:
            raise ValueError(
                f"Formato de archivo no soportado: {filename}. "
                "Solo se permiten archivos .txt, .xlsx, .docx"
            )

    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        logger.error(f"Error procesando archivo {filename}: {str(e)}")
        raise ValueError(f"Error al procesar el archivo: {str(e)}")

    if not identification_numbers:
        raise ValueError(
            "No se encontraron números de identificación válidos en el archivo"
        )

    # Convertir a lista y ordenar
    result = sorted(list(identification_numbers))
    logger.info(f"Se encontraron {len(result)} números de identificación únicos")

    return result


def _validate_identification_number(number_str: str) -> bool:
    """
    Valida que un string sea un número de identificación válido.

    :param number_str: String a validar
    :return: True si es válido, False en caso contrario
    """
    # Remover espacios y caracteres especiales comunes
    cleaned = number_str.strip().replace(",", "").replace(".", "")

    # Verificar que solo contenga dígitos
    if not cleaned.isdigit():
        return False

    # Verificar longitud (cédulas colombianas: 6-10 dígitos)
    if len(cleaned) < 6 or len(cleaned) > 10:
        return False

    # Verificar que no sea un número obviamente inválido
    if cleaned == "0" * len(cleaned):  # Todos ceros
        return False

    return True
